# -*- coding: utf-8 -*-
# MIT License
#
# Copyright (c) 2025 A B Vijay Kumar
#
# Permission is hereby granted, free of charge, to any person obtaining a copy
# of this software and associated documentation files (the "Software"), to deal
# in the Software without restriction, including without limitation the rights
# to use, copy, modify, merge, publish, distribute, sublicense, and/or sell
# copies of the Software, and to permit persons to whom the Software is
# furnished to do so, subject to the following conditions:
#
# The above copyright notice and this permission notice shall be included in all
# copies or substantial portions of the Software.
#
# THE SOFTWARE IS PROVIDED "AS IS", WITHOUT WARRANTY OF ANY KIND, EXPRESS OR
# IMPLIED, INCLUDING BUT NOT LIMITED TO THE WARRANTIES OF MERCHANTABILITY,
# FITNESS FOR A PARTICULAR PURPOSE AND NONINFRINGEMENT. IN NO EVENT SHALL THE
# AUTHORS OR COPYRIGHT HOLDERS BE LIABLE FOR ANY CLAIM, DAMAGES OR OTHER
# LIABILITY, WHETHER IN AN ACTION OF CONTRACT, TORT OR OTHERWISE, ARISING FROM,
# OUT OF OR IN CONNECTION WITH THE SOFTWARE OR THE USE OR OTHER DEALINGS IN THE
# SOFTWARE.

import re
import os
import sys
import glob
import argparse
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
import markdown
from bs4 import BeautifulSoup
from playwright.sync_api import sync_playwright

class MarkdownToWordConverter:
    def __init__(self):
        self.log_callback = None
    
    def log(self, message):
        """Log message using callback if available"""
        if self.log_callback:
            self.log_callback(message)
        else:
            print(message)
    
    def add_hyperlink(self, paragraph, url, text):
        """Add a hyperlink to a paragraph"""
        # Create hyperlink
        part = paragraph.part
        r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

        # Create the w:hyperlink tag and add needed values
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)

        # Create a new run object
        new_run = OxmlElement('w:r')

        # Create a new run properties tag and add color and underline
        rPr = OxmlElement('w:rPr')

        # Set color property to blue
        c = OxmlElement('w:color')
        c.set(qn('w:val'), "0000FF")
        rPr.append(c)

        # Set underline property
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)

        new_run.append(rPr)

        # Create text element
        new_run.text = text

        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)

        return hyperlink
    
    def process_text_with_formatting(self, element, paragraph):
        """Process text with inline formatting like bold, italic, links"""
        if element is None:
            return
            
        # Handle different types of content
        if hasattr(element, 'contents'):
            for content in element.contents:
                if hasattr(content, 'name') and content.name:
                    if content.name == 'strong' or content.name == 'b':
                        run = paragraph.add_run(content.get_text())
                        run.bold = True
                    elif content.name == 'em' or content.name == 'i':
                        run = paragraph.add_run(content.get_text())
                        run.italic = True
                    elif content.name == 'code':
                        run = paragraph.add_run(content.get_text())
                        run.font.name = 'Courier New'
                        run.font.size = Pt(10)
                    elif content.name == 'a':
                        href = content.get('href', '')
                        text = content.get_text()
                        if href:
                            self.add_hyperlink(paragraph, href, text)
                        else:
                            paragraph.add_run(text)
                    else:
                        # For other tags, process recursively or just add text
                        self.process_text_with_formatting(content, paragraph)
                else:
                    # Plain text node
                    paragraph.add_run(str(content))
        else:
            # Element doesn't have contents, just get its text
            paragraph.add_run(element.get_text() if hasattr(element, 'get_text') else str(element))
    
    def process_list(self, doc, list_element, level=0):
        """Process ordered and unordered lists"""
        for li in list_element.find_all('li', recursive=False):
            # Create paragraph for list item
            p = doc.add_paragraph()
            
            # Add appropriate bullet or number
            if list_element.name == 'ul':
                # Unordered list - add bullet
                if level == 0:
                    bullet = '• '
                elif level == 1:
                    bullet = '○ '
                else:
                    bullet = '▪ '
                p.add_run(bullet)
            else:  # ol
                # For ordered lists, we'll use a simple numbering
                # In a more advanced implementation, you'd track the number
                p.add_run('1. ')
            
            # Add indentation for nested lists
            p.paragraph_format.left_indent = Inches(0.5 * (level + 1))
            
            # Process the content of the list item
            # Check if there are nested lists
            nested_lists = li.find_all(['ul', 'ol'], recursive=False)
            
            # Get text content excluding nested lists
            li_copy = BeautifulSoup(str(li), 'html.parser')
            for nested_list in li_copy.find_all(['ul', 'ol']):
                nested_list.decompose()
            
            # Add the text content
            li_element = li_copy.find('li')
            if li_element:
                self.process_text_with_formatting(li_element, p)
            
            # Process nested lists
            for nested_list in nested_lists:
                self.process_list(doc, nested_list, level + 1)
    
    def process_table(self, doc, table_element):
        """Process HTML table and convert to Word table"""
        rows = table_element.find_all('tr')
        if not rows:
            return
            
        # Count columns from the first row
        first_row = rows[0]
        cols = len(first_row.find_all(['td', 'th']))
        
        if cols == 0:
            return
            
        # Create Word table
        table = doc.add_table(rows=len(rows), cols=cols)
        table.style = 'Table Grid'
        
        for row_idx, row in enumerate(rows):
            cells = row.find_all(['td', 'th'])
            for col_idx, cell in enumerate(cells):
                if col_idx < len(table.rows[row_idx].cells):
                    table.rows[row_idx].cells[col_idx].text = cell.get_text().strip()
                    # Make header row bold
                    if cell.name == 'th':
                        for paragraph in table.rows[row_idx].cells[col_idx].paragraphs:
                            for run in paragraph.runs:
                                run.bold = True
        
    def extract_mermaid_diagrams(self, markdown_text):
        """Extract mermaid diagrams from markdown text"""
        pattern = r'```mermaid\n(.*?)\n```'
        diagrams = re.findall(pattern, markdown_text, re.DOTALL)
        return diagrams
    
    def save_mermaid_files(self, diagrams, chapter_number, output_dir='.'):
        """Save mermaid diagrams as .mmd files"""
        mmd_files = []
        for i, diagram in enumerate(diagrams, 1):
            filename = "ch{}_{}".format(chapter_number, i) + ".mmd"
            filepath = os.path.join(output_dir, filename)
            
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write(diagram.strip())
            
            mmd_files.append(filepath)
            print("Created mermaid file: " + filepath)
        
        return mmd_files
    
    def render_mermaid_to_image(self, mermaid_code, output_path):
        """Render mermaid diagram to image using playwright"""
        self.log(f"  Rendering mermaid diagram: {os.path.basename(output_path)}")
        html_content = '''
        <!DOCTYPE html>
        <html>
        <head>
            <script src="https://cdn.jsdelivr.net/npm/mermaid/dist/mermaid.min.js"></script>
        </head>
        <body>
            <div class="mermaid">''' + mermaid_code + '''</div>
            <script>
                mermaid.initialize({startOnLoad: true});
            </script>
        </body>
        </html>
        '''
        
        with sync_playwright() as p:
            browser = p.chromium.launch()
            page = browser.new_page()
            page.set_content(html_content)
            page.wait_for_selector('.mermaid svg')
            page.locator('.mermaid').screenshot(path=output_path)
            browser.close()
        self.log(f"  ✓ Diagram rendered successfully")
    
    def convert_file(self, md_file, output_file, chapter_number=1):
        """Convert single markdown file to Word document"""
        self.log(f"Reading file: {os.path.basename(md_file)}")
        with open(md_file, 'r', encoding='utf-8') as f:
            markdown_text = f.read()
        
        doc = Document()
        
        # Extract and render mermaid diagrams
        mermaid_diagrams = self.extract_mermaid_diagrams(markdown_text)
        if mermaid_diagrams:
            self.log(f"  Found {len(mermaid_diagrams)} mermaid diagram(s)")
        diagram_images = []
        
        # Save mermaid diagrams as .mmd files if any exist
        if mermaid_diagrams:
            output_dir = os.path.dirname(output_file) if os.path.dirname(output_file) else '.'
            self.save_mermaid_files(mermaid_diagrams, chapter_number, output_dir)
        
        for i, diagram in enumerate(mermaid_diagrams):
            temp_image = "temp_diagram_" + str(i) + ".png"
            self.render_mermaid_to_image(diagram, temp_image)
            diagram_images.append(temp_image)
        
        # Remove mermaid blocks from markdown
        markdown_text = re.sub(r'```mermaid\n.*?\n```', '{{MERMAID_PLACEHOLDER}}', markdown_text, flags=re.DOTALL)
        
        # Convert markdown to HTML with extensions
        self.log(f"  Converting markdown to Word format...")
        html = markdown.markdown(markdown_text, extensions=['extra', 'codehilite', 'toc'])
        soup = BeautifulSoup(html, 'html.parser')
        
        # Process HTML elements
        diagram_index = 0
        for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'ul', 'ol', 'blockquote', 'pre', 'table', 'hr']):
            if '{{MERMAID_PLACEHOLDER}}' in element.get_text():
                if diagram_index < len(diagram_images):
                    doc.add_picture(diagram_images[diagram_index], width=Inches(6))
                    diagram_index += 1
            elif element.name.startswith('h') and element.name[1:].isdigit():
                level = int(element.name[1])
                doc.add_heading(element.get_text(), level)
            elif element.name in ['ul', 'ol']:
                self.process_list(doc, element)
            elif element.name == 'blockquote':
                p = doc.add_paragraph()
                self.process_text_with_formatting(element, p)
                p.paragraph_format.left_indent = Inches(0.5)
                p.paragraph_format.right_indent = Inches(0.5)
            elif element.name == 'pre':
                # Code block
                code_text = element.get_text()
                p = doc.add_paragraph(code_text)
                p.style = 'Normal'
                for run in p.runs:
                    run.font.name = 'Courier New'
                    run.font.size = Pt(10)
                p.paragraph_format.left_indent = Inches(0.5)
            elif element.name == 'table':
                self.process_table(doc, element)
            elif element.name == 'hr':
                # Horizontal rule - add a paragraph with text
                p = doc.add_paragraph('─' * 50)
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            else:
                # Regular paragraph
                p = doc.add_paragraph()
                self.process_text_with_formatting(element, p)
        
        # Ensure output directory exists
        output_dir = os.path.dirname(output_file)
        if output_dir:
            os.makedirs(output_dir, exist_ok=True)
        
        # Save document
        self.log(f"  Saving document: {os.path.basename(output_file)}")
        doc.save(output_file)
        
        # Cleanup temp files
        for img in diagram_images:
            if os.path.exists(img):
                os.remove(img)
        self.log(f"  ✓ Conversion complete")
    
    def convert_combined(self, md_files, output_file):
        """Convert multiple markdown files into a single Word document"""
        self.log(f"Creating combined document from {len(md_files)} files")
        doc = Document()
        
        for chapter_idx, md_file in enumerate(md_files, 1):
            self.log(f"Processing file {chapter_idx}/{len(md_files)}: {os.path.basename(md_file)}")
            with open(md_file, 'r', encoding='utf-8') as f:
                markdown_text = f.read()
            
            # Add file title
            filename = os.path.splitext(os.path.basename(md_file))[0]
            doc.add_heading(filename, 1)
            
            # Extract and render mermaid diagrams
            mermaid_diagrams = self.extract_mermaid_diagrams(markdown_text)
            if mermaid_diagrams:
                self.log(f"  Found {len(mermaid_diagrams)} mermaid diagram(s)")
            diagram_images = []
            
            # Save mermaid diagrams as .mmd files if any exist
            if mermaid_diagrams:
                output_dir = os.path.dirname(output_file) if os.path.dirname(output_file) else '.'
                self.save_mermaid_files(mermaid_diagrams, chapter_idx, output_dir)
            
            for i, diagram in enumerate(mermaid_diagrams):
                temp_image = "temp_diagram_" + filename + "_" + str(i) + ".png"
                self.render_mermaid_to_image(diagram, temp_image)
                diagram_images.append(temp_image)
            
            # Remove mermaid blocks from markdown
            markdown_text = re.sub(r'```mermaid\n.*?\n```', '{{MERMAID_PLACEHOLDER}}', markdown_text, flags=re.DOTALL)
            
            # Convert markdown to HTML with extensions
            self.log(f"  Converting content to Word format...")
            html = markdown.markdown(markdown_text, extensions=['extra', 'codehilite', 'toc'])
            soup = BeautifulSoup(html, 'html.parser')
            
            # Process HTML elements
            diagram_index = 0
            for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'ul', 'ol', 'blockquote', 'pre', 'table', 'hr']):
                if '{{MERMAID_PLACEHOLDER}}' in element.get_text():
                    if diagram_index < len(diagram_images):
                        doc.add_picture(diagram_images[diagram_index], width=Inches(6))
                        diagram_index += 1
                elif element.name.startswith('h') and element.name[1:].isdigit():
                    level = int(element.name[1]) + 1  # Offset by 1 since file title is H1
                    doc.add_heading(element.get_text(), level)
                elif element.name in ['ul', 'ol']:
                    self.process_list(doc, element)
                elif element.name == 'blockquote':
                    p = doc.add_paragraph()
                    self.process_text_with_formatting(element, p)
                    p.paragraph_format.left_indent = Inches(0.5)
                    p.paragraph_format.right_indent = Inches(0.5)
                elif element.name == 'pre':
                    # Code block
                    code_text = element.get_text()
                    p = doc.add_paragraph(code_text)
                    p.style = 'Normal'
                    for run in p.runs:
                        run.font.name = 'Courier New'
                        run.font.size = Pt(10)
                    p.paragraph_format.left_indent = Inches(0.5)
                elif element.name == 'table':
                    self.process_table(doc, element)
                elif element.name == 'hr':
                    # Horizontal rule - add a paragraph with text
                    p = doc.add_paragraph('─' * 50)
                    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                else:
                    # Regular paragraph
                    p = doc.add_paragraph()
                    self.process_text_with_formatting(element, p)
            
            # Cleanup temp files for this file
            for img in diagram_images:
                if os.path.exists(img):
                    os.remove(img)
            
            # Add page break between files (except for the last file)
            if md_file != md_files[-1]:
                doc.add_page_break()
            
            self.log(f"  ✓ File {chapter_idx}/{len(md_files)} processed")
        
        # Ensure output directory exists
        output_dir = os.path.dirname(output_file)
        if output_dir:
            os.makedirs(output_dir, exist_ok=True)
        
        # Save combined document
        self.log(f"Saving combined document: {os.path.basename(output_file)}")
        doc.save(output_file)
        self.log(f"✓ Combined document created successfully")
    
    def convert_folder(self, folder_path, target_dir=None, combine=False):
        """Convert markdown files in folder to Word documents"""
        md_files = sorted(glob.glob(os.path.join(folder_path, '*.md')))
        
        if not md_files:
            print("No markdown files found in " + folder_path)
            return
        
        # Set up target directory
        if target_dir is None:
            target_dir = os.path.join(folder_path, 'docx')
        
        # Create target directory if it doesn't exist
        os.makedirs(target_dir, exist_ok=True)
        
        if combine:
            output_file = os.path.join(target_dir, "combined.docx")
            self.convert_combined(md_files, output_file)
            print("Converted: " + output_file)
        else:
            for chapter_idx, md_file in enumerate(md_files, 1):
                filename = os.path.splitext(os.path.basename(md_file))[0]
                output_file = os.path.join(target_dir, filename + ".docx")
                self.convert_file(md_file, output_file, chapter_idx)
                print("Converted: " + output_file)

def main():
    parser = argparse.ArgumentParser(description='Convert markdown files to Word documents with mermaid diagram support')
    parser.add_argument('folder_path', help='Path to folder containing markdown files')
    parser.add_argument('-c', '--combine', action='store_true', help='Combine all markdown files into a single Word document')
    parser.add_argument('-t', '--target-dir', help='Target directory to store output docx files (default: creates "docx" folder in source folder)')
    
    parser.epilog = '''
    Mermaid Diagrams:
    When mermaid diagrams are found, they are automatically:
    1. Rendered as images and embedded in the Word document
    2. Saved as separate .mmd files with naming: ch<chapter>_<diagram>.mmd
    
    Example: The 3rd diagram in the 2nd chapter becomes "ch2_3.mmd"
    '''
    
    args = parser.parse_args()
    
    converter = MarkdownToWordConverter()
    converter.convert_folder(args.folder_path, args.target_dir, args.combine)
    converter.log("Conversion completed for all files")

if __name__ == "__main__":
    main()
